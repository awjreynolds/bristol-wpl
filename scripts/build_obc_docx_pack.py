#!/usr/bin/env python3
from __future__ import annotations

import csv
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACK_DIR = ROOT / "business_case" / "obc" / "docx-pack"
RENDER_DIR = ROOT / "business_case" / "obc" / "docx-pack-render-qa"
OBC_MD = ROOT / "business_case" / "obc" / "simulation-release" / "bristol-wpl-outline-business-case-simulation-release.md"

OUTPUTS = {
    "obc": PACK_DIR / "bristol-wpl-obc-simulation-release.docx",
    "guide": PACK_DIR / "bristol-wpl-obc-reader-support-guide.docx",
    "risk": PACK_DIR / "bristol-wpl-obc-risk-process-control-summary.docx",
}
ZIP_PATH = PACK_DIR / "bristol-wpl-obc-document-pack.zip"
MANIFEST_PATH = PACK_DIR / "PACK-MANIFEST.txt"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(17, 24, 39)
MUTED = RGBColor(75, 85, 99)
LIGHT_FILL = "F2F4F7"
NOTICE_FILL = "FFF4CC"


def clean_text(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text.strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def configure_doc(doc: Document, running_label: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = running_label
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header.runs:
        set_run_font(header.runs[0], size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Simulation-only document pack. Not an approved OBC, officer advice, consultation material or statutory submission."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if footer.runs:
        set_run_font(footer.runs[0], size=8.5, color=MUTED)


def add_title_page(doc: Document, title: str, subtitle: str, doc_type: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Bristol Workplace Parking Levy")
    set_run_font(run, size=11, color=MUTED, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=24, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(subtitle)
    set_run_font(run, size=13, color=MUTED)

    for label, value in [
        ("Document type", doc_type),
        ("Status", "Simulation-only editable DOCX"),
        ("Date", "2026-06-27"),
        ("Pack stage", "Stage 35A DOCX document pack"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)

    add_notice(
        doc,
        "Read this first",
        "This document is part of a shareable DOCX pack generated from the public Bristol WPL simulation repository. It is not a Bristol City Council document, not legal advice, not officer advice, not consultation material, not procurement authority and not an approved OBC.",
    )


def add_notice(doc: Document, title: str, body: str, fill: str = NOTICE_FILL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, size=10.5, color=INK, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_inline_markup(p, body)
    doc.add_paragraph()


def add_inline_markup(paragraph, text: str) -> None:
    # Supports simple **bold** markers and strips Markdown code ticks for Word readers.
    parts = re.split(r"(\*\*[^*]+\*\*)", text.replace("`", ""))
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def add_para(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    add_inline_markup(p, text)


def add_bullet(doc: Document, text: str) -> None:
    add_para(doc, clean_text(text), style="List Bullet")


def add_number(doc: Document, text: str) -> None:
    add_para(doc, clean_text(text), style="List Number")


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [clean_text(cell) for cell in line.split("|")]


def is_separator_row(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    table = doc.add_table(rows=len(padded), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(padded):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_markup(p, value)
            for run in p.runs:
                if row_index == 0:
                    run.bold = True
                run.font.size = Pt(9.5 if width > 3 else 10)
        if row_index == 0:
            set_repeat_table_header(table.rows[0])
    doc.add_paragraph()


def add_markdown_file(doc: Document, path: Path, skip_initial_h1: bool = False) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    index = 0
    skipped_h1 = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line.strip():
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                if not is_separator_row(lines[index]):
                    table_lines.append(parse_table_row(lines[index]))
                index += 1
            add_table(doc, table_lines)
            continue
        heading_match = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_text(heading_match.group(2))
            if skip_initial_h1 and level == 1 and skipped_h1 < 2:
                skipped_h1 += 1
                index += 1
                continue
            doc.add_heading(text, level=min(level, 3))
            index += 1
            continue
        if line.startswith("- "):
            add_bullet(doc, line[2:])
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.*)$", line)
        if numbered:
            add_number(doc, numbered.group(1))
            index += 1
            continue
        add_para(doc, line)
        index += 1


def read_csv(rel: str) -> list[dict[str, str]]:
    with (ROOT / rel).open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def add_rows_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    add_table(doc, [headers] + rows)


def open_items(rows: list[dict[str, str]], status_column: str = "status") -> list[dict[str, str]]:
    return [row for row in rows if row.get(status_column) not in {"closed", "complete"}]


def build_obc_docx() -> None:
    doc = Document()
    configure_doc(doc, "Bristol WPL OBC simulation release")
    add_title_page(
        doc,
        "OBC Simulation Release",
        "Editable Word version generated from the Stage 33A Markdown release",
        "Outline Business Case simulation",
    )
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_markdown_file(doc, OBC_MD, skip_initial_h1=True)
    doc.save(OUTPUTS["obc"])


def build_reader_support_docx() -> None:
    doc = Document()
    configure_doc(doc, "Bristol WPL OBC reader support guide")
    add_title_page(
        doc,
        "OBC Reader Support Guide",
        "How to read, share and challenge the simulation document pack",
        "Reading support guide",
    )

    doc.add_heading("What this pack is", level=1)
    add_para(doc, "This pack is a Word document set for people who do not want to navigate a developer-style repository. It packages the simulation OBC and the main reading controls into shareable DOCX files.")
    add_para(doc, "It does not approve a Bristol Workplace Parking Levy, create officer advice, launch consultation, authorise procurement, settle the statutory route or replace legal, finance, modelling, equalities, consultation, data or commercial review.")

    doc.add_heading("Documents in the ZIP", level=1)
    add_rows_table(
        doc,
        ["Document", "Purpose", "How to use it"],
        [
            ["bristol-wpl-obc-simulation-release.docx", "The OBC-shaped simulation document.", "Read for the draft business-case structure and the evidence gaps that prevent reliance."],
            ["bristol-wpl-obc-reader-support-guide.docx", "This guide.", "Use first if you are an officer, cabinet member, legal reviewer or non-technical reader."],
            ["bristol-wpl-obc-risk-process-control-summary.docx", "Risk, issue, process and gate summary.", "Use to understand what the simulation found, which controls exist and what remains blocked."],
        ],
    )

    doc.add_heading("Suggested reading route", level=1)
    for item in [
        "Read the first page of this guide.",
        "Open the OBC simulation release and read the Notice and Release Use sections.",
        "Review the Executive Summary and Conclusions sections.",
        "Open the risk/process/control summary and check the P0/P1 blockers.",
        "Use the repo only for detailed source paths, registers and audit trail.",
    ]:
        add_number(doc, item)

    doc.add_heading("Safe statements and prohibited inferences", level=1)
    add_rows_table(
        doc,
        ["Safe to say", "Do not infer"],
        [
            ["A complete editable OBC simulation exists.", "That Bristol City Council, WECA/MCA, DfT or Secretary of State has approved it."],
            ["The simulation records risks, issues, evidence gaps and controls.", "That the risk ratings, mitigations or evidence base are professionally assured."],
            ["The OBC is useful for learning, critique and future drafting.", "That it can be used as officer advice, consultation material or procurement authority."],
            ["DOCX files are provided for officer-friendly sharing.", "That the DOCX pack is an official distribution pack or publication pack."],
        ],
    )

    doc.add_heading("Where the detailed trail lives", level=1)
    add_rows_table(
        doc,
        ["Need", "Repository location"],
        [
            ["Risk register", "governance/risk_register.csv"],
            ["Issues and blockers", "governance/issues_register.csv"],
            ["Evidence gaps", "evidence/evidence_gap_register.csv"],
            ["Pitfalls and lessons", "governance/pitfalls_register.csv"],
            ["Stage-gate process", "docs/stages/README.md and review/stage_gate_reports/"],
            ["OBC controls", "business_case/obc/controls/"],
        ],
    )
    doc.save(OUTPUTS["guide"])


def build_risk_process_docx() -> None:
    doc = Document()
    configure_doc(doc, "Bristol WPL OBC risk process control summary")
    add_title_page(
        doc,
        "OBC Risk, Process And Control Summary",
        "Support document for the Stage 33A OBC simulation release and Stage 35A DOCX pack",
        "Risk and process support summary",
    )

    doc.add_heading("Current gate position", level=1)
    add_notice(
        doc,
        "Current no-go position",
        "The simulation remains no-go for approval, consultation, OBC reliance, FBC reliance, procurement authority and statutory submission. The DOCX pack changes format only; it does not close any evidence or authority gap.",
    )

    issues = open_items(read_csv("governance/issues_register.csv"))
    priority_issues = [row for row in issues if row.get("severity") in {"P0", "P1"}][:14]
    doc.add_heading("Priority blockers", level=1)
    add_rows_table(
        doc,
        ["ID", "Stage", "Severity", "Issue"],
        [[row["issue_id"], row["stage"], row["severity"], row["issue"]] for row in priority_issues],
    )

    risks = open_items(read_csv("governance/risk_register.csv"))
    priority_risks = [row for row in risks if row.get("gross_rating") in {"P0", "P1"}][:14]
    doc.add_heading("Priority risks", level=1)
    add_rows_table(
        doc,
        ["ID", "Stage", "Rating", "Risk", "Residual"],
        [[row["risk_id"], row["stage"], row["gross_rating"], row["risk"], row["residual_rating"]] for row in priority_risks],
    )

    gaps = open_items(read_csv("evidence/evidence_gap_register.csv"))
    material_gaps = [row for row in gaps if row.get("materiality") == "high"][:14]
    doc.add_heading("High-materiality evidence gaps", level=1)
    add_rows_table(
        doc,
        ["ID", "Stage", "Gap", "Status"],
        [[row["gap_id"], row["stage"], row["gap"], row["status"]] for row in material_gaps],
    )

    doc.add_heading("OBC gate checklist summary", level=1)
    checklist = read_csv("business_case/obc/controls/stage-7-obc-gate-checklist.csv")
    add_rows_table(
        doc,
        ["Gate item", "Assurance area", "Status", "Gate effect"],
        [[row["gate_item_id"], row["assurance_area"], row["current_status"], row["gate_effect"]] for row in checklist],
    )

    doc.add_page_break()
    doc.add_heading("Process trail", level=1)
    add_rows_table(
        doc,
        ["Stage", "Purpose", "Main output"],
        [
            ["Stage 32A", "WECA-style OBC simulation drafting.", "Simulated working draft and exemplar corpus."],
            ["Stage 33A", "Release the editable OBC simulation separately from the blocked assembled path.", "Simulation release Markdown and gate report."],
            ["Stage 34A", "Adopt a bounded GOV.UK-style writing skill.", "Public/officer writing-control skill and no-overclaim checks."],
            ["Stage 35A", "Create an officer-friendly DOCX pack.", "DOCX files and ZIP package for sharing."],
        ],
    )

    doc.add_heading("Next proof needed before real reliance", level=1)
    for item in [
        "Source-supported Bristol final order-maker, submitter and signatory route.",
        "Current-law WECA/MCA role and funding-dependency disposition.",
        "WPL-specific DfT process or logged DfT response.",
        "Authoritative boundary, workplace parking inventory and displacement evidence.",
        "OAR, ASR, ASST, model outputs, BCR and VFM category.",
        "Consultation launch authority, materials, privacy, accessibility and response-analysis route.",
        "Real legal, finance, modelling, equalities, data, commercial, consultation and officer review.",
    ]:
        add_bullet(doc, item)

    doc.save(OUTPUTS["risk"])


def write_manifest() -> None:
    manifest = """Bristol WPL OBC DOCX document pack
Status: Stage 35A simulation-only officer-friendly distribution pack
Date: 2026-06-27

Contents:
- bristol-wpl-obc-simulation-release.docx
- bristol-wpl-obc-reader-support-guide.docx
- bristol-wpl-obc-risk-process-control-summary.docx

Limits:
- Not an approved Bristol OBC.
- Not officer advice.
- Not legal advice.
- Not consultation material.
- Not procurement authority.
- Not a WECA/MCA, DfT or Secretary of State submission.
- Not for real-world reliance.

Source repository paths:
- business_case/obc/simulation-release/bristol-wpl-outline-business-case-simulation-release.md
- governance/risk_register.csv
- governance/issues_register.csv
- evidence/evidence_gap_register.csv
- business_case/obc/controls/
"""
    MANIFEST_PATH.write_text(manifest, encoding="utf-8")


def write_zip() -> None:
    with zipfile.ZipFile(ZIP_PATH, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in OUTPUTS.values():
            archive.write(path, arcname=path.name)
        archive.write(MANIFEST_PATH, arcname=MANIFEST_PATH.name)


def main() -> int:
    PACK_DIR.mkdir(parents=True, exist_ok=True)
    build_obc_docx()
    build_reader_support_docx()
    build_risk_process_docx()
    write_manifest()
    write_zip()
    print(f"Wrote DOCX pack to {PACK_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
